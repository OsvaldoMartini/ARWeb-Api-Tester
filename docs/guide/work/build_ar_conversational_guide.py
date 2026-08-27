"""Build the complete client-facing AR Conversational Word manual."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_arapi_guide as base
from ar_conversational_content import (
    AGENTS,
    API_ROUTES,
    CONTROL_INDEX,
    GLOSSARY,
    NAV_ITEMS,
    PROVIDERS,
    REVIEW_DATE,
    SCREENS,
    TOC,
    TROUBLESHOOTING,
    VERSION,
)


WORK_DIR = Path(__file__).resolve().parent
GUIDE_DIR = WORK_DIR.parent
SCREEN_DIR = GUIDE_DIR / "screenshots-ar-conversational"
OUTPUT_PATH = GUIDE_DIR / "AR-Conversational-Complete-Client-Guide.docx"


def configure_document(doc: Document) -> None:
    """Apply the approved compact-reference/editorial-cover presentation."""
    base.configure_document(doc)
    section = doc.sections[0]

    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    left = p.add_run("AR Conversational | Complete Client Guide")
    base.set_run_font(left, size=9, color=base.MUTED, bold=True)
    right = p.add_run(f"\tVersion {VERSION}")
    base.set_run_font(right, size=9, color=base.MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    label = fp.add_run("AR Conversational  |  ")
    base.set_run_font(label, size=9, color=base.MUTED)
    base.add_field(fp, "PAGE")

    core = doc.core_properties
    core.title = "AR Conversational Complete Client Guide"
    core.subject = "Client user manual for AR Conversational 1.0.0"
    core.author = "AR Conversational Documentation Team"
    core.last_modified_by = "AR Conversational Documentation Team"
    core.keywords = "AR Conversational, banking assistant, ARAPI, client guide"
    core.comments = (
        "Reviewed against the AR Conversational 1.0.0 desktop source and "
        "captured application screens."
    )


def build_cover(doc: Document) -> None:
    top = doc.add_paragraph()
    base.bookmark_paragraph(top, "doc_top")
    top.paragraph_format.space_after = Pt(0)
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("CLIENT USER MANUAL")
    base.set_run_font(run, size=11, color=base.BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AR Conversational")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Complete Client Guide")

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run("Ask. Route. Review. Hand off.")
    base.set_run_font(run, size=11, color=base.MUTED, italic=True)
    tag.paragraph_format.space_after = Pt(70)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run(f"Version {VERSION}  |  Windows Desktop")
    base.set_run_font(run, size=11, color=base.NAVY, bold=True)
    version.paragraph_format.space_after = Pt(4)

    reviewed = doc.add_paragraph()
    reviewed.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = reviewed.add_run(f"Reviewed {REVIEW_DATE}")
    base.set_run_font(run, size=9.5, color=base.MUTED)
    base.page_break(doc)


def build_front_matter(doc: Document) -> None:
    base.add_heading(doc, "Document control", 1, "document_control")
    base.add_table(
        doc,
        ["Field", "Value"],
        [
            ("Product", "AR Conversational Windows desktop application"),
            ("Guide version", VERSION),
            ("Audience", "Bank employees, client-demo operators, implementation teams, support, and administrators"),
            ("Scope", "AR Conversational; ARAPI authoring and execution are referenced only at the handoff boundary"),
            ("Evidence basis", "Version 1.0.0 source review, successful frontend build, running Tauri desktop application, and reproduced screenshots"),
            ("Review date", REVIEW_DATE),
        ],
        [2700, 6660],
        9.4,
    )
    base.add_callout(
        doc,
        "Client-safe scope",
        "Screens use local demonstration data and contain no API keys. Never enter production secrets, personal data, real account identifiers, or confidential documents into a demo or support session.",
        "caution",
    )
    base.add_heading(doc, "How to use this guide", 2)
    base.add_list(
        doc,
        [
            "Read Sections 2-4 before operating a new installation or client demonstration.",
            "Use Section 6 for the visual walkthrough; every figure comes from the running AR Conversational desktop app.",
            "Use Sections 7-9 to understand modes, automatic routing, explicit agent selection, evidence, and answer limitations.",
            "Use Section 10 when conversational instructions cross into ARAPI catalog, BotJob, or execution work.",
            "Use Appendix A to locate every visible control and Appendix E for client acceptance.",
        ],
    )
    base.page_break(doc)

    toc = base.add_heading(doc, "Contents", 1, "toc")
    toc.paragraph_format.space_after = Pt(10)
    for anchor, label in TOC:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15 if not label.startswith("Appendix") else 0.35)
        p.paragraph_format.space_after = Pt(4)
        base.add_internal_link(
            p,
            label,
            anchor,
            size=10.5,
            bold=not label.startswith("Appendix"),
        )
    base.add_callout(
        doc,
        "Navigation",
        "In Word, hold Ctrl while clicking an entry. In most PDF viewers, click normally. Each major section includes a Back to contents link.",
        "note",
    )
    base.page_break(doc)


def build_about(doc: Document) -> None:
    base.add_heading(doc, "1. About AR Conversational", 1, "about", True)
    base.add_para(
        doc,
        "AR Conversational is a Windows desktop banking-assistant interface for employee and client-demonstration conversations. It routes questions to named banking agents, relates answers to the API catalog shared with ARAPI, and delegates catalog/test/BotJob instructions to the ARAPI assistant workflow.",
    )
    base.add_para(
        doc,
        "The React interface communicates with a localhost C# backend on port 8787. AR Conversational and ARAPI share catalog and settings data. The conversational app intentionally keeps endpoint import, workflow editing, generated scripts, and execution review in the separate ARAPI application.",
    )
    base.add_heading(doc, "Core concepts", 2)
    for label, text in (
        ("Conversation mode", "Employee or Client context used for automatic routing candidates, examples, input wording, and supported answer presentation."),
        ("Banking agent", "A named routing profile with a mode, purpose, and banking keywords."),
        ("Auto-route", "The default selection that chooses among agents for the current mode by keyword score."),
        ("Explicit agent", "A manually selected agent that overrides automatic mode filtering for the next questions."),
        ("Evidence", "Up to three method/path records selected from the shared catalog for a normal banking question."),
        ("Delegation", "Handoff of catalog/test/BotJob wording to the ARAPI assistant route."),
        ("Synthetic answer", "A local demonstration response produced from built-in mock records, never live core-banking data."),
    ):
        base.add_definition(doc, label, text)

    base.add_heading(doc, "Application boundary", 2)
    base.add_table(
        doc,
        ["Use AR Conversational for", "Use ARAPI for"],
        [
            ("Employee/client banking questions", "OpenAPI/Swagger import and catalog ownership"),
            ("Automatic or explicit banking-agent routing", "BotJob creation and detailed workflow editing"),
            ("Viewing catalog evidence and limitations", "Bash/curl generation and controlled command-line execution"),
            ("Read-only endpoint search", "Test execution history, reports, environments, and mock controls"),
            ("Shared AI-provider settings", "Reviewing or approving delegated assistant actions"),
        ],
        [4680, 4680],
        9.2,
    )


def build_release(doc: Document) -> None:
    base.add_heading(doc, "2. Read this before using the current release", 1, "release", True)
    base.add_callout(
        doc,
        "Version basis",
        f"This manual describes the observed AR Conversational {VERSION} desktop build reviewed on {REVIEW_DATE}. Catalog counts, provider state, and demo records will differ by workstation.",
        "note",
    )
    base.add_heading(doc, "Current-build qualifications", 2)
    base.add_list(
        doc,
        [
            "Balance/account answers in the reviewed build use local synthetic customer accounts with masked IBANs. They are not live or core-banking data and must not support a customer or financial decision.",
            "Visible conversation turns are held in the mounted page only. Navigating away, reloading, or restarting clears them; version 1.0.0 has no conversation-history workspace.",
            "Auto-route filters candidate agents by Employee/Client mode and scores keywords. A manually selected agent overrides that automatic mode filter.",
            "Endpoint evidence is a keyword-based catalog match, not API-contract validation, authorization proof, or runtime-success evidence.",
            "Normal /agents/ask and delegated /app-assistant/chat answers are deterministic in the reviewed C# path. They do not prove that the selected external provider generated the text.",
            "Provider Test reports a local configuration indicator in the reviewed build. Connected/OK and a zero-millisecond value are not an external-provider round-trip.",
            "Provider records and the selected default are shared with ARAPI. A change in either application affects the shared setting.",
            "Provider-list responses omit stored keys and the UI masks entry. The reviewed save path does not independently demonstrate an encryption transform, so protect the local database as credential-bearing.",
            "Delegated run wording can execute all saved BotJobs against the built-in Mock environment. Always inspect execution history in ARAPI and never interpret a local green result as a real banking-system change.",
            "Port 5174 is the development Vite UI port. The packaged Tauri desktop app loads bundled frontend files and does not require a public listener on that port.",
        ],
    )
    base.add_callout(
        doc,
        "Production safeguard",
        "Treat AR Conversational as a local assistant and demonstration interface. Human review, customer authorization, source-system verification, and the organization's banking controls remain mandatory.",
        "risk",
    )


def build_install(doc: Document) -> None:
    base.add_heading(doc, "3. Installation and first launch", 1, "install", True)
    base.add_heading(doc, "Install or unpack", 2)
    base.add_list(
        doc,
        [
            "Use the AR Conversational installer supplied by your organization, or unpack the approved portable bundle into a user-writable directory.",
            "Keep the executable, packaged backend sidecar, and any supplied data directory together. Do not move individual binary files out of the distribution.",
            "Launch AR Conversational. A resizable desktop window opens; no application login is shown in version 1.0.0.",
        ],
        numbered=True,
    )
    base.add_heading(doc, "Verify first launch", 2)
    base.add_list(
        doc,
        [
            "Confirm the top-right sidecar indicator becomes online.",
            "Open Home and confirm Shared backend is online on port 8787.",
            "Review Agent catalog coverage. A partial count means some agents currently have no keyword-matching catalog endpoints; it does not mean the installation failed.",
            "Open API Catalog and confirm expected imported endpoints are visible. If empty, import specifications in ARAPI.",
            "Open Settings and confirm local service information and the intended shared default provider record.",
        ],
    )
    base.add_heading(doc, "Shared local data", 2)
    base.add_para(
        doc,
        "Both desktop applications use the same local database selected by their shell. A portable distribution may use data\\arweb.db beside the executable when a data directory is supplied; other bundles use an application-data location. Back up the exact deployment path before upgrades or migrations.",
    )
    base.add_callout(
        doc,
        "Local-only services",
        "The backend is intended for localhost access. Do not bind, proxy, or firewall-publish port 8787 to an untrusted network.",
        "caution",
    )


def build_quick_start(doc: Document) -> None:
    base.add_heading(doc, "4. End-to-end quick start", 1, "quick_start", True)
    base.add_list(
        doc,
        [
            "Start AR Conversational and wait for sidecar online.",
            "On first use, choose Start as Bank Employee or Start as e-Banking Client. Choose the perspective appropriate for the demonstration; this is not user authentication.",
            "Open Coverage and review the currently wired agents. Close Coverage when finished.",
            "Leave the agent selector on Auto-route for the first question, or choose a named agent when the responsibility is known.",
            "Enter one specific, non-sensitive question and choose Send (or press Enter). Do not include real customer or credential data.",
            "Read the complete answer and every limitation. Expand endpoints used when evidence is available, then validate the method/path in API Catalog or ARAPI.",
            "If the response is delegated to ARAPI, open ARAPI and review the catalog match, BotJob, mock execution, or other action before acceptance.",
            "Capture only approved, sanitized evidence. Remember that leaving the page clears visible conversation turns.",
        ],
        numbered=True,
    )
    base.add_heading(doc, "Safe first questions", 2)
    base.add_code_block(
        doc,
        "Employee: Which catalog endpoints support account-balance servicing?\n"
        "Client demo: Show my demonstration account balances.\n"
        "Catalog handoff: Search the catalog for payment endpoints.",
    )
    base.add_callout(
        doc,
        "Never paste",
        "Real names, IBANs, account numbers, authentication tokens, API keys, confidential messages, production payloads, or regulated customer documents.",
        "risk",
    )


def build_interface(doc: Document) -> None:
    base.add_heading(doc, "5. Interface fundamentals", 1, "interface", True)
    base.add_heading(doc, "Global layout", 2)
    base.add_definition(doc, "Sidebar", "Four destinations: Home, AR Conversational, API Catalog, and Settings.")
    base.add_definition(doc, "Page title", "The upper-left top bar repeats the current workspace name.")
    base.add_definition(doc, "Employee / Client", "The top-bar mode buttons change the in-memory conversation perspective. They do not authenticate a user or change the catalog.")
    base.add_definition(doc, "sidecar indicator", "Shows whether the React UI can reach the localhost backend. It refreshes approximately every five seconds.")
    base.add_definition(doc, "No login required", "A product-state notice, not a security guarantee. Windows access and data-file permissions form part of the boundary.")
    base.add_heading(doc, "Sidebar destinations", 2)
    base.add_table(doc, ["Destination", "Purpose"], NAV_ITEMS, [2400, 6960], 9.4)
    base.add_heading(doc, "Common interaction states", 2)
    base.add_list(
        doc,
        [
            "Blue filled button: selected mode or main page action.",
            "Outlined button: available secondary action, such as Mode, Coverage, Test, or Set default.",
            "Dimmed button: disabled until required text or configuration exists, or represents the current default.",
            "Green indicator: online, configured default, or Connected according to its label; read the release qualification for Provider Test.",
            "Orange limitation panel: mandatory warning about synthetic data, evidence quality, or another constraint.",
            "Chevron: expands or collapses provider details or endpoint evidence.",
        ],
    )
    base.add_heading(doc, "Keyboard and session behavior", 2)
    base.add_table(
        doc,
        ["Action", "Behavior"],
        [
            ("Enter in the question field", "Submits non-empty text when no request is already running."),
            ("Mouse/keyboard activate Send", "Sends one local request and appends the result to this mounted page."),
            ("Navigate to another sidebar page", "Unmounts the conversation workspace and clears its visible turns."),
            ("Reload/restart", "Clears visible turns and starts a new UI session."),
            ("Mode change", "Changes routing context for subsequent questions; it is not identity or authorization."),
        ],
        [2900, 6460],
        9.2,
    )


def build_screens(doc: Document) -> None:
    base.add_heading(doc, "6. Screen-by-screen guide", 1, "screens", True)
    base.add_para(
        doc,
        "The following 15 figures show every main AR Conversational screen and the important interactive states captured from the running desktop application. The values are demonstration data; the controls and safeguards are the focus.",
    )
    for index, screen in enumerate(SCREENS, start=1):
        if index > 1:
            base.page_break(doc)
        filename, title, alt, note = screen
        base.add_heading(doc, f"6.{index} {title}", 2)
        base.add_figure(doc, filename, title, alt, note)

        if filename == "01-home.png":
            base.add_callout(doc, "Coverage reading", "13/14 in this capture is dynamic keyword coverage, not a fixed product limit or a successful API-call count.", "note")
        elif filename == "02-assistant-employee.png":
            base.add_callout(doc, "Session scope", "Questions and answers remain visible only while this conversation page stays mounted.", "caution")
        elif filename == "03-mode-welcome.png":
            base.add_callout(doc, "Mode is not identity", "Start as e-Banking Client changes routing and wording. It does not authenticate a customer or authorize account access.", "risk")
        elif filename == "04-agent-coverage.png":
            base.add_callout(doc, "Match quality", "A wired label/count indicates keyword overlap with catalog records. Validate each endpoint's real contract and ownership in ARAPI.", "caution")
        elif filename == "05-agent-selected.png":
            base.add_callout(doc, "Override", "Relationship Manager remains forced until Auto-route or another agent is selected, even if the active mode would normally choose another group.", "note")
        elif filename == "06-assistant-client.png":
            base.add_callout(doc, "Client-safe demonstration", "Use synthetic questions only. The client mode is a presentation context, not a connected e-banking session.", "risk")
        elif filename == "07-client-balance-answer.png":
            base.add_callout(doc, "Synthetic balances", "All shown account records are local demonstration data with masked identifiers. Never quote them as a customer position.", "risk")
        elif filename == "08-client-evidence-expanded.png":
            base.add_callout(doc, "Evidence boundary", "Method/path pills are catalog matches. They do not show that authorization, request validation, network execution, or a source-system response occurred.", "caution")
        elif filename == "09-arapi-delegation.png":
            base.add_callout(doc, "Human handoff", "Open ARAPI and review any matched endpoint, synthetic record, BotJob, or run. The conversational card is not final approval.", "caution")
        elif filename == "10-api-catalog.png":
            base.add_callout(doc, "Read-only", "Import, edit, and mapping controls intentionally remain in ARAPI.", "note")
        elif filename == "11-catalog-filtered.png":
            base.add_callout(doc, "Search semantics", "Filtering uses method/path/summary text. A broad word such as balance may include technically or semantically unsuitable endpoints.", "caution")
        elif filename == "12-settings.png":
            base.add_callout(doc, "Shared setting", "Changing provider records or default selection here changes the shared setting seen by ARAPI.", "caution")
        elif filename == "13-settings-provider-form.png":
            base.add_callout(doc, "Credential handling", "Blank after a saved key preserves it. Protect the local data store and never capture a typed production key.", "risk")
        elif filename == "14-settings-local-services.png":
            base.add_callout(doc, "Port scope", "8787 is the localhost backend; 5174 is development UI only. Neither should be advertised as a public client service.", "caution")
        elif filename == "15-settings-test-result.png":
            base.add_callout(doc, "Connected qualification", "In the reviewed build, this proves only that a local provider record exists. It does not prove external network, credentials, model entitlement, or LLM processing.", "risk")


def build_modes(doc: Document) -> None:
    base.add_heading(doc, "7. Employee and client modes", 1, "modes", True)
    base.add_para(
        doc,
        "Mode controls the conversation perspective used by the interface and automatic agent routing. It is not login, customer identity, entitlement, consent, or upstream API authorization.",
    )
    base.add_table(
        doc,
        ["Behavior", "Employee mode", "Client mode"],
        [
            ("Automatic candidates", "Nine employee agents", "Five client agents"),
            ("Example language", "Servicing, operations, portfolio, controls", "Balances, payments, trading, lending, messages"),
            ("Input presentation", "Bank-employee question context", "e-Banking client-style question context"),
            ("Synthetic balance wording", "Employee-oriented explanation", "Client-oriented account summary"),
            ("Authentication effect", "None", "None"),
            ("Catalog effect", "None; shared inventory", "None; shared inventory"),
        ],
        [2350, 3505, 3505],
        8.9,
    )
    base.add_heading(doc, "Choose or change mode", 2)
    base.add_list(
        doc,
        [
            "At first use, read the welcome comparison and select Start as Bank Employee or Start as e-Banking Client.",
            "Later, select the top-bar Employee/Client button for a direct change, or select Mode in the workspace to reopen the comparison.",
            "Ask a new question after the change. Previously displayed text is not reprocessed or relabeled.",
            "Return the agent selector to Auto-route if you want mode-filtered automatic routing; an explicit agent selection still takes precedence.",
        ],
        numbered=True,
    )
    base.add_heading(doc, "First-use modal behavior", 2)
    base.add_para(
        doc,
        "The introduction is remembered locally after closing or choosing a mode. Use the Mode button to reopen it. Clearing browser/webview storage may cause it to appear again.",
    )
    base.add_callout(doc, "Client demonstrations", "State explicitly that Client mode is synthetic and local before showing balances or messages. Do not let the visual context imply connection to a customer's bank session.", "risk")


def build_agents(doc: Document) -> None:
    base.add_heading(doc, "8. Banking agents and routing", 1, "agents", True)
    base.add_heading(doc, "How Auto-route works", 2)
    base.add_list(
        doc,
        [
            "The backend takes the current Employee/Client mode and builds that mode's candidate list.",
            "It compares banking keywords from the question with each candidate's routing profile.",
            "The highest-scoring candidate handles the normal banking question; deterministic fallback behavior applies when scoring is weak.",
            "If a named agent is selected, that explicit ID is used instead of the mode-filtered automatic selection.",
            "The answer returns the agent name plus any evidence and limitations; inspect them before relying on the response.",
        ],
        numbered=True,
    )
    base.add_callout(doc, "Routing is classification", "Agent choice does not grant data access, verify customer identity, validate an API contract, or invoke a live banking system.", "caution")
    base.add_heading(doc, "Agent groups", 2)
    base.add_table(
        doc,
        ["Mode", "Agent", "Primary purpose"],
        [(mode, name, purpose) for mode, name, purpose, _ in AGENTS],
        [1250, 2860, 5250],
        8.7,
    )
    base.add_heading(doc, "Coverage panel", 2)
    base.add_para(
        doc,
        "Coverage calls the local agent/capability inventory and displays current keyword-match counts against the shared catalog. Counts change after ARAPI imports or catalog changes. Zero means no current keyword match; it is not an agent permission state.",
    )
    base.add_heading(doc, "When to select an agent explicitly", 2)
    base.add_list(
        doc,
        [
            "The organizational owner is known and the question is intentionally narrow.",
            "You are testing a specific agent's wording or evidence selection.",
            "Auto-route repeatedly chooses an adjacent responsibility for ambiguous wording.",
            "You are demonstrating how an agent behaves outside its normal mode. Label that as an override and not normal routing.",
        ],
    )
    base.add_callout(doc, "Reset after testing", "Choose Auto-route when explicit-agent testing is finished, especially before changing from Employee to Client mode.", "success")


def build_conversations(doc: Document) -> None:
    base.add_heading(doc, "9. Questions, answers, limitations, and evidence", 1, "conversations", True)
    base.add_heading(doc, "Write an effective question", 2)
    base.add_list(
        doc,
        [
            "Ask one business question or instruction at a time.",
            "Use specific banking nouns and actions: account balance, payment status, portfolio holdings, loan repayment, compliance screening, settlement, or statement.",
            "State whether you want an explanation, a catalog search, or an ARAPI test/BotJob handoff.",
            "Use synthetic identifiers only; remove all production credentials and customer data.",
            "If evidence is important, ask for the relevant endpoint area and then validate it in API Catalog/ARAPI.",
        ],
    )
    base.add_heading(doc, "Request lifecycle", 2)
    base.add_table(
        doc,
        ["Stage", "What happens", "What the operator must verify"],
        [
            ("1. Submit", "UI sends question, mode, and optional agent to localhost.", "No sensitive text; intended mode/agent."),
            ("2. Classify", "Normal banking question or delegated ARAPI intent is selected.", "The returned agent/action matches the request."),
            ("3. Answer", "Deterministic local answer or delegated action result is returned.", "Read the complete content; do not infer external AI use."),
            ("4. Evidence", "Up to three catalog records may be attached.", "Method/path/summary and real contract suitability."),
            ("5. Limitations", "Synthetic/no-match/other warnings may be shown.", "Treat every limitation as mandatory acceptance context."),
            ("6. Handoff", "ARAPI review may be required.", "Open the separate app and inspect saved artifacts/history."),
        ],
        [1150, 4120, 4090],
        8.5,
    )
    base.add_heading(doc, "Synthetic balance response", 2)
    base.add_para(
        doc,
        "Balance/account wording is the one normal-question area with a detailed local demonstration response in the reviewed backend. It reads built-in synthetic account rows, masks identifiers, totals the mock values, and appends a clear limitation. It does not query a core-banking system.",
    )
    base.add_callout(doc, "Financial-use prohibition", "Do not communicate synthetic balances to a customer, reconcile them, make an investment/credit/payment decision from them, or retain them as evidence of a real position.", "risk")
    base.add_heading(doc, "Evidence interpretation", 2)
    base.add_definition(doc, "Method/path", "The endpoint record selected from the imported catalog.")
    base.add_definition(doc, "Why it appeared", "Keyword overlap between the expanded question and endpoint metadata.")
    base.add_definition(doc, "What it proves", "Only that a catalog record was selected for display.")
    base.add_definition(doc, "What it does not prove", "Schema correctness, data entitlement, authentication, network execution, response validity, or successful business effect.")
    base.add_heading(doc, "Conversation retention", 2)
    base.add_para(
        doc,
        "Version 1.0.0 does not persist visible turns as a searchable conversation history. If the organization needs an audit record, capture an approved sanitized summary and the corresponding ARAPI/source-system evidence before navigating away. Do not use screenshots as the only live-system proof.",
    )


def build_delegation(doc: Document) -> None:
    base.add_heading(doc, "10. ARAPI delegation and cross-application workflow", 1, "delegation", True)
    base.add_para(
        doc,
        "Prompts containing test creation/building, BotJob, catalog-search, catalog-for, or payment-endpoint intent are sent to the shared ARAPI assistant route rather than the normal banking-agent route.",
    )
    base.add_heading(doc, "Common delegated intents", 2)
    base.add_table(
        doc,
        ["Intent", "Typical wording", "Observed local result"],
        [
            ("Search/list", "Search the catalog for payment endpoints", "Returns matched endpoint records and a handoff warning."),
            ("Create/build", "Create a BotJob for an account flow", "May use an imported match or create a synthetic mock-only record and saved BotJob."),
            ("Run", "Run the BotJobs", "Can execute all saved BotJobs against the built-in Mock environment and return local summaries."),
        ],
        [1700, 3440, 4220],
        8.8,
    )
    base.add_heading(doc, "Required handoff procedure", 2)
    base.add_list(
        doc,
        [
            "Read the conversational action/result card and identify whether it searched, created, or ran anything.",
            "Open the separate ARAPI application using the same local deployment/database.",
            "For search: validate method, path, summary, imported source, banking mapping, and ownership.",
            "For create: open BotJob Designer and review endpoint origin, every command, variable, body, header, ordering, enabled state, and environment.",
            "For run: open Execute Tests/History and identify every job that ran, the Mock target, status rows, and limitations.",
            "Delete or quarantine unwanted synthetic/demo artifacts only through approved ARAPI data-management procedures.",
            "Obtain human approval before generating Bash/curl or pointing any environment at a non-mock target.",
        ],
        numbered=True,
    )
    base.add_callout(doc, "No silent acceptance", "A delegated card is an instruction result, not approval of an API contract, BotJob, test outcome, or production change.", "risk")
    base.add_heading(doc, "Recommended division of responsibilities", 2)
    base.add_table(
        doc,
        ["Role", "AR Conversational responsibility", "ARAPI responsibility"],
        [
            ("Business user", "Ask a sanitized question; read limitations.", "Review business intent and expected effects."),
            ("API owner", "Interpret catalog evidence cautiously.", "Confirm real specification, method/path, schemas, and authorization."),
            ("QA/operator", "Identify the requested handoff.", "Review BotJob, environment, generated script, run, and target evidence."),
            ("Administrator", "Protect workstation/local service.", "Protect shared DB, credentials, backups, and deployment bundle."),
        ],
        [1750, 3650, 3960],
        8.6,
    )


def build_catalog(doc: Document) -> None:
    base.add_heading(doc, "11. Read-only API Catalog", 1, "catalog", True)
    base.add_para(
        doc,
        "API Catalog displays endpoints imported by ARAPI into the shared database. AR Conversational does not import, edit, delete, or map specifications.",
    )
    base.add_heading(doc, "Search", 2)
    base.add_list(
        doc,
        [
            "Open API Catalog from the sidebar or Home shortcut.",
            "Enter a method, path fragment, or summary word in Search method / path.",
            "Review every returned row: HTTP method, path, summary, and mapping status.",
            "Clear or replace the search term to restore the broader inventory.",
            "Open ARAPI for source/import details, corrections, or workflow selection.",
        ],
        numbered=True,
    )
    base.add_heading(doc, "Interpret the rows", 2)
    base.add_table(
        doc,
        ["Field", "Meaning", "Review point"],
        [
            ("Method", "GET, POST, PUT, DELETE, and other imported verbs.", "Write methods may have side effects in a real target."),
            ("Path", "Imported route template.", "Confirm parameter names and base URL in the source contract."),
            ("Summary", "Specification description used by search/routing.", "May be incomplete or ambiguous."),
            ("Mapping", "Banking taxonomy status from ARAPI.", "Unmapped does not automatically mean unusable; requires owner review."),
        ],
        [1500, 3900, 3960],
        8.9,
    )
    base.add_callout(doc, "Catalog stewardship", "When results are missing, stale, duplicated, or incorrectly described, correct the trusted specification/import in ARAPI. Do not try to compensate with confidential prompt text.", "caution")


def build_providers(doc: Document) -> None:
    base.add_heading(doc, "12. AI provider settings", 1, "providers", True)
    base.add_para(
        doc,
        "Settings stores provider records shared by ARAPI and AR Conversational. The current deterministic conversation routes do not prove that the selected provider generated an answer; configure these records only for approved shared features and future integration needs.",
    )
    base.add_heading(doc, "Supported records", 2)
    base.add_table(
        doc,
        ["Provider", "Default model shown", "Base URL field"],
        PROVIDERS,
        [2600, 3900, 2860],
        8.9,
    )
    base.add_heading(doc, "Configure or update", 2)
    base.add_list(
        doc,
        [
            "Expand the intended provider card.",
            "Confirm or enter Model. For Azure, Ollama, or Custom, confirm Base URL according to the approved deployment.",
            "Enter a scoped, revocable API key only when required. The field masks typed characters.",
            "Select Save. After a key is stored, leaving the key field blank on a later save preserves it.",
            "Select SET DEFAULT on a configured provider when it should be the shared default. If exactly one keyed provider exists, the backend may promote it automatically.",
            "Select Test and interpret the result using the qualification below; perform any approved real provider/network validation separately.",
        ],
        numbered=True,
    )
    base.add_heading(doc, "Badges and controls", 2)
    base.add_table(
        doc,
        ["Control/state", "Meaning"],
        [
            ("NO KEY", "No stored key is reported for the provider; default selection is disabled."),
            ("SET DEFAULT", "Available for a keyed, non-default provider; changes the shared default."),
            ("DEFAULT", "The provider is the current shared default; choose another configured provider to change it."),
            ("Test", "Returns the reviewed build's local provider-record indicator."),
            ("Connected / OK / 0 ms", "Does not prove an external request, credential validity, model access, or provider-generated response."),
        ],
        [2800, 6560],
        9.0,
    )
    base.add_heading(doc, "Credential safeguards", 2)
    base.add_list(
        doc,
        [
            "Use non-production or narrowly scoped keys wherever possible; rotate them according to policy.",
            "Do not put keys in screenshots, manuals, chat questions, logs, or support attachments.",
            "Provider-list API responses omit key values and the UI masks input, but the local database must still be classified and protected as credential-bearing.",
            "Restrict Windows account access and data-directory ACLs; back up and dispose of the database under credential-handling policy.",
            "Before claiming provider connectivity, perform an approved end-to-end check that actually contacts the provider and records sanitized evidence.",
        ],
    )
    base.add_callout(doc, "UI label qualification", "Although the Settings page presents an encryption-at-rest statement, the reviewed C# save path does not independently demonstrate an encryption transform. Do not rely on that label as the only credential control.", "risk")


def build_security(doc: Document) -> None:
    base.add_heading(doc, "13. Security and operational practices", 1, "security", True)
    base.add_heading(doc, "Security boundary", 2)
    base.add_para(
        doc,
        "AR Conversational has no internal login in version 1.0.0. The managed Windows account, endpoint controls, local file permissions, backend localhost binding, API credentials, source-system authorization, and organizational procedures form the operational boundary.",
    )
    base.add_heading(doc, "Required practices", 2)
    base.add_list(
        doc,
        [
            "Install only an approved, signed distribution on an organization-managed Windows device.",
            "Keep backend port 8787 on localhost and block unapproved network exposure.",
            "Protect and back up the shared local database; it may contain catalog data, BotJobs, provider records, and credential-bearing fields.",
            "Use synthetic demonstration data. Do not enter production personal, banking, authentication, or confidential information into questions.",
            "Treat screenshots, reports, provider state, catalog details, and delegated ARAPI artifacts according to their data classification.",
            "Require human review and source-system evidence for any real API action or customer-facing statement.",
            "Separate development, mock, UAT, and production responsibilities; never infer target authorization from the selected conversation mode.",
            "Record application version, operator, sanitized question/purpose, agent/mode, ARAPI handoff, evidence location, and exceptions for controlled testing.",
        ],
    )
    base.add_heading(doc, "What not to enter", 2)
    base.add_table(
        doc,
        ["Prohibited input", "Use instead"],
        [
            ("Real customer name, IBAN, account number", "Synthetic names and clearly fake identifiers."),
            ("API key, OAuth token, password, session cookie", "Approved secret store or runtime environment in ARAPI workflows."),
            ("Confidential payment/portfolio/credit content", "Sanitized scenario that preserves only the test intent."),
            ("Production request/response body", "Schema-only example with fabricated values."),
            ("Regulated documents or messages", "Approved redacted sample designed for testing."),
        ],
        [4550, 4810],
        9.0,
    )
    base.add_heading(doc, "Operational evidence", 2)
    base.add_para(
        doc,
        "For real-system acceptance, combine the sanitized conversational request with reviewed ARAPI artifacts and authoritative target evidence: approved request/response capture, service logs, audit record, database/state confirmation, or a business-owner reconciliation. A conversational answer or screenshot alone is insufficient.",
    )
    base.add_callout(doc, "No-login does not mean no security", "Do not deploy the app on a shared/unmanaged account or assume Client mode identifies a customer. Apply workstation, file, credential, network, and source-system controls.", "caution")


def build_troubleshooting(doc: Document) -> None:
    base.add_heading(doc, "14. Troubleshooting", 1, "troubleshooting", True)
    base.add_para(
        doc,
        "When escalating, include AR Conversational version, installation type, exact displayed message, current mode/agent, sanitized question, backend status, and reproduction steps. Remove credentials and customer data.",
    )
    for issue, resolution in TROUBLESHOOTING:
        base.add_heading(doc, issue, 2)
        base.add_para(doc, resolution)
    base.add_heading(doc, "Safe diagnostic information", 2)
    base.add_list(
        doc,
        [
            "AR Conversational version and installer/portable distribution type.",
            "Home sidecar status, backend port, and dynamic agent-coverage count.",
            "Current mode and selected agent name (or Auto-route).",
            "A fully sanitized reproduction question and the displayed limitation text.",
            "Endpoint methods/paths shown as evidence when those catalog details are approved for sharing.",
            "Provider name/model/base-URL hostname and badge state; never the key or Authorization data.",
            "Whether ARAPI shows the same catalog/provider data and whether a delegated artifact/run exists.",
        ],
    )


def build_status(doc: Document) -> None:
    base.add_heading(doc, "15. Feature status matrix", 1, "status", True)
    rows = [
        ("Employee and Client modes", "Available", "Presentation/routing context only; not authentication."),
        ("14 named banking agents", "Available", "Nine employee and five client routing profiles."),
        ("Auto-route", "Available", "Mode filter plus keyword score; explicit agent overrides."),
        ("Agent coverage", "Available", "Dynamic catalog keyword counts, not execution proof."),
        ("Normal banking answers", "Available with qualification", "Deterministic local response and catalog evidence."),
        ("Synthetic balance answer", "Demo only", "Local mock accounts; no core-banking connection."),
        ("Endpoint evidence", "Available with qualification", "Up to three keyword matches; validate contract in ARAPI."),
        ("Conversation history", "Not persisted", "Visible turns clear on navigation/reload/restart."),
        ("ARAPI search/list delegation", "Available", "Returns catalog matches for human review."),
        ("ARAPI BotJob creation delegation", "Available with qualification", "May create synthetic mock-only endpoint/job."),
        ("ARAPI run delegation", "Available with qualification", "Can run all saved jobs against built-in Mock."),
        ("API Catalog", "Read-only", "Search method/path/summary; import/edit remains in ARAPI."),
        ("Seven provider records", "Available", "Shared with ARAPI; protect the local data store."),
        ("Provider default selection", "Available", "Shared record state; not proof of LLM use."),
        ("Provider Test", "Local indicator", "Not an external provider round-trip in reviewed build."),
        ("Frontend development port 5174", "Development only", "Not required as a public listener in packaged desktop."),
        ("Backend port 8787", "Available", "Localhost-only shared C# backend."),
    ]
    base.add_table(doc, ["Capability", "Status", "Qualification"], rows, [3200, 2200, 3960], 8.5)


def build_appendices(doc: Document) -> None:
    base.page_break(doc)
    base.add_heading(doc, "Appendix A. Complete control index", 1, "controls", True)
    base.add_para(
        doc,
        "This index covers all visible navigation, buttons, selectors, fields, disclosure controls, status controls, and important conditional actions in AR Conversational 1.0.0.",
    )
    base.add_table(
        doc,
        ["Area", "Control", "Available when", "Result"],
        CONTROL_INDEX,
        [1350, 2240, 1900, 3870],
        8.0,
    )

    base.page_break(doc)
    base.add_heading(doc, "Appendix B. Agent reference", 1, "agent_reference", True)
    base.add_para(
        doc,
        "Keywords summarize routing intent; the live score uses backend profiles and the wording in the current question. Coverage counts depend on the imported catalog.",
    )
    base.add_table(
        doc,
        ["Mode", "Agent", "Purpose", "Example routing terms"],
        AGENTS,
        [1050, 2360, 3670, 2280],
        7.8,
    )

    base.page_break(doc)
    base.add_heading(doc, "Appendix C. Local API reference", 1, "api", True)
    base.add_callout(
        doc,
        "Administrator/developer reference",
        "These are the localhost JSON routes used by the AR Conversational UI. Normal client use should stay in the desktop interface. Never publish the API publicly or send it secrets outside approved settings workflows.",
        "caution",
    )
    base.add_table(doc, ["Method", "Route", "Purpose"], API_ROUTES, [900, 4050, 4410], 8.3)

    base.page_break(doc)
    base.add_heading(doc, "Appendix D. Glossary", 1, "glossary", True)
    base.add_table(doc, ["Term", "Meaning"], GLOSSARY, [2500, 6860], 9.2)

    base.page_break(doc)
    base.add_heading(doc, "Appendix E. Client acceptance checklist", 1, "handoff", True)
    base.add_heading(doc, "Application and security", 2)
    base.add_list(
        doc,
        [
            "Version and distribution type are recorded; the approved bundle launches without a global Node.js dependency.",
            "sidecar online appears after a clean launch and port 8787 remains localhost-only.",
            "The exact shared local database path, backup/restore procedure, and ACL owner are documented.",
            "Operators understand that there is no internal login and Client mode is not authentication.",
        ],
    )
    base.add_heading(doc, "Conversation and catalog", 2)
    base.add_list(
        doc,
        [
            "Employee and Client modes, Auto-route, and explicit agent selection behave as documented.",
            "All 14 agents are listed and coverage changes when the shared catalog changes.",
            "A synthetic balance demonstration displays masked identifiers and a visible no-live-data limitation.",
            "Evidence can be expanded and is understood as keyword matching, not execution proof.",
            "API Catalog search is read-only and the ARAPI import/correction owner is identified.",
            "Conversation turns are confirmed to clear on navigation/reload/restart.",
        ],
    )
    base.add_heading(doc, "ARAPI handoff", 2)
    base.add_list(
        doc,
        [
            "A catalog-search instruction produces a reviewable ARAPI delegation card.",
            "Created/synthetic BotJobs are reviewed in ARAPI before save, script generation, or execution.",
            "Any delegated run is reconciled against ARAPI History and identified as built-in Mock unless separately proven.",
            "Human approval and target-side evidence are required for non-mock acceptance.",
        ],
    )
    base.add_heading(doc, "Providers and support", 2)
    base.add_list(
        doc,
        [
            "Provider records/default selection are confirmed as shared with ARAPI.",
            "Keys are scoped, revocable, absent from screenshots, and the local database is protected as credential-bearing.",
            "Provider Test is accepted only as a local configuration indicator; any external validation is performed separately.",
            "Support procedures list sanitized evidence and explicitly prohibit credentials/customer data in tickets.",
        ],
    )
    base.add_callout(
        doc,
        "Sign-off",
        "Record approver, date, application build, workstation/deployment, shared-data location, tested modes/agents, catalog source, evidence location, exceptions, and remediation owner.",
        "success",
    )
    bottom = doc.add_paragraph()
    base.bookmark_paragraph(bottom, "doc_bottom")
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.add_internal_link(bottom, "Return to contents", "toc", size=9.5)


def build() -> Path:
    base.SCREEN_DIR = SCREEN_DIR
    base._figure_number = 0
    base._bookmark_id = 1
    base._next_num_id = 100

    doc = Document()
    configure_document(doc)
    base.install_numbering(doc)
    build_cover(doc)
    build_front_matter(doc)
    build_about(doc)
    build_release(doc)
    build_install(doc)
    build_quick_start(doc)
    build_interface(doc)
    build_screens(doc)
    build_modes(doc)
    build_agents(doc)
    build_conversations(doc)
    build_delegation(doc)
    build_catalog(doc)
    build_providers(doc)
    build_security(doc)
    build_troubleshooting(doc)
    build_status(doc)
    build_appendices(doc)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
