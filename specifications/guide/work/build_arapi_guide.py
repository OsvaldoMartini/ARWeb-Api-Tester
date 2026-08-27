"""Build the complete client-facing ARAPI Word manual."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from guide_content import (
    API_ROUTES,
    CATEGORIES,
    COMMANDS,
    CONTROL_INDEX,
    GLOSSARY,
    NAV_ITEMS,
    REVIEW_DATE,
    SCREENS,
    TOC,
    TROUBLESHOOTING,
    VERSION,
)


WORK_DIR = Path(__file__).resolve().parent
GUIDE_DIR = WORK_DIR.parent
SCREEN_DIR = GUIDE_DIR / "screenshots"
OUTPUT_PATH = GUIDE_DIR / "ARAPI-Complete-Client-Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
INK = "202A35"
MUTED = "5F6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_GOLD = "FFF4D6"
GOLD = "7A5A00"
PALE_RED = "FDECEC"
RED = "9B1C1C"
PALE_GREEN = "E9F7EF"
GREEN = "166534"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, font: str, size: float, color: str = INK,
                    bold: bool = False, before: float = 0, after: float = 6,
                    line: float = 1.25, keep_next: bool = False) -> None:
    style.font.name = font
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    pf.widow_control = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    configure_style(doc.styles["Normal"], "Calibri", 11, INK, False, 0, 6, 1.25)
    configure_style(doc.styles["Heading 1"], "Calibri", 16, BLUE, True, 18, 10, 1.0, True)
    configure_style(doc.styles["Heading 2"], "Calibri", 13, BLUE, True, 14, 7, 1.0, True)
    configure_style(doc.styles["Heading 3"], "Calibri", 12, DARK_BLUE, True, 10, 5, 1.0, True)
    configure_style(doc.styles["Title"], "Calibri", 30, NAVY, True, 0, 8, 1.0, True)
    configure_style(doc.styles["Subtitle"], "Calibri", 15, "2B5163", False, 0, 4, 1.0, True)
    configure_style(doc.styles["Caption"], "Calibri", 9, MUTED, False, 2, 8, 1.0, True)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    left = p.add_run("ARAPI | Complete Client Guide")
    set_run_font(left, size=9, color=MUTED, bold=True)
    right = p.add_run(f"\tVersion {VERSION}")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    label = fp.add_run("ARAPI  |  ")
    set_run_font(label, size=9, color=MUTED)
    add_field(fp, "PAGE")

    core = doc.core_properties
    core.title = "ARAPI Complete Client Guide"
    core.subject = "Client user manual for ARAPI 1.0.0"
    core.author = "ARAPI Documentation Team"
    core.last_modified_by = "ARAPI Documentation Team"
    core.keywords = "ARAPI, API testing, BotJob, Bash, curl, client guide"
    core.comments = "Reviewed against the ARAPI 1.0.0 desktop source and captured application screens."


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


_bookmark_id = 1


def bookmark_paragraph(paragraph, name: str) -> None:
    global _bookmark_id
    safe = "".join(ch if ch.isalnum() or ch == "_" else "_" for ch in name)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(_bookmark_id))
    start.set(qn("w:name"), safe[:40])
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(_bookmark_id))
    paragraph._p.append(start)
    paragraph._p.append(end)
    _bookmark_id += 1


def add_internal_link(paragraph, text: str, anchor: str, color: str = BLUE,
                      bold: bool = False, size: float = 10) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_external_link(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_heading(doc: Document, text: str, level: int, bookmark: str | None = None,
                back_to_toc: bool = False):
    p = doc.add_heading(text, level=level)
    if bookmark:
        bookmark_paragraph(p, bookmark)
    if back_to_toc:
        nav = doc.add_paragraph()
        nav.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        nav.paragraph_format.space_after = Pt(2)
        # The section heading already keeps with this navigation paragraph.
        # Keeping the navigation paragraph with the next block prevents the
        # heading/link pair from being orphaned at the foot of a page.
        nav.paragraph_format.keep_with_next = True
        add_internal_link(nav, "Back to contents", "toc", size=8.5)
    return p


def add_para(doc: Document, text: str = "", *, bold_lead: str | None = None,
             italic: bool = False, color: str | None = None,
             align=WD_ALIGN_PARAGRAPH.LEFT, after: float | None = None):
    p = doc.add_paragraph()
    p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, bold=True, color=color or INK)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic, color=color or INK)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic, color=color or INK)
    return p


def add_callout(doc: Document, label: str, text: str, kind: str = "note") -> None:
    palette = {
        "note": (LIGHT_BLUE, DARK_BLUE),
        "caution": (PALE_GOLD, GOLD),
        "risk": (PALE_RED, RED),
        "success": (PALE_GREEN, GREEN),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    # This one-row callout container needs an explicit first-row semantic for
    # assistive technology even though it is primarily used for visual layout.
    mark_repeat_header(table.rows[0])
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    lr = p.add_run(f"{label}: ")
    set_run_font(lr, size=10, color=accent, bold=True)
    tr = p.add_run(text)
    set_run_font(tr, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    # Code samples use a one-row table to preserve shading and wrapping.
    mark_repeat_header(table.rows[0])
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, 110, 110, 150, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    lines = code.rstrip().splitlines() or [""]
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color="17212B")
        if index < len(lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]],
              widths_dxa: Sequence[int], font_size: float = 9.2):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    mark_repeat_header(hdr)
    prevent_row_split(hdr)
    for index, header in enumerate(headers):
        cell = hdr.cells[index]
        cell.text = ""
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(header)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


_abstract_bullet_id = 50
_abstract_decimal_id = 51
_next_num_id = 100


def install_numbering(doc: Document) -> None:
    numbering = doc.part.numbering_part.element
    for abstract_id, fmt, text, left, hanging in (
        (_abstract_bullet_id, "bullet", "\u2022", 540, 270),
        (_abstract_decimal_id, "decimal", "%1.", 540, 270),
    ):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)


def new_num_id(doc: Document, abstract_id: int) -> int:
    global _next_num_id
    num_id = _next_num_id
    _next_num_id += 1
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    # Word can otherwise infer that adjacent lists using the same abstract
    # definition are continuations. An explicit level override makes every
    # procedure created by add_list() restart at 1.
    if abstract_id == _abstract_decimal_id:
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        num.append(level_override)
    numbering.append(num)
    return num_id


def add_list(doc: Document, items: Sequence[str], numbered: bool = False) -> None:
    num_id = new_num_id(doc, _abstract_decimal_id if numbered else _abstract_bullet_id)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num)
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_definition(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.18)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=DARK_BLUE)
    body = p.add_run(text)
    set_run_font(body, color=INK)


_figure_number = 0


def add_figure(doc: Document, filename: str, title: str, alt: str, note: str) -> None:
    global _figure_number
    _figure_number += 1
    path = SCREEN_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(6.25))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt)
    doc_pr.set("title", title)
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(f"Figure {_figure_number}. {title}")
    set_run_font(cap_run, size=9, color=MUTED, italic=True)
    bookmark_paragraph(caption, f"fig{_figure_number}")
    add_para(doc, note, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_cover(doc: Document) -> None:
    top = doc.add_paragraph()
    bookmark_paragraph(top, "doc_top")
    top.paragraph_format.space_after = Pt(0)
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("CLIENT USER MANUAL")
    set_run_font(kr, size=11, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("ARAPI")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Complete Client Guide")
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tag.add_run("Import. Design. Execute. Export.")
    set_run_font(tr, size=11, color=MUTED, italic=True)
    tag.paragraph_format.space_after = Pt(70)
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = version.add_run(f"Version {VERSION}  |  Windows Desktop")
    set_run_font(vr, size=11, color=NAVY, bold=True)
    version.paragraph_format.space_after = Pt(4)
    reviewed = doc.add_paragraph()
    reviewed.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = reviewed.add_run(f"Reviewed {REVIEW_DATE}")
    set_run_font(rr, size=9.5, color=MUTED)
    page_break(doc)


def build_front_matter(doc: Document) -> None:
    add_heading(doc, "Document control", 1, "document_control")
    add_table(
        doc,
        ["Field", "Value"],
        [
            ("Product", "ARAPI Windows desktop application"),
            ("Guide version", VERSION),
            ("Audience", "Banking operations staff, QA analysts, implementation teams, and administrators"),
            ("Scope", "ARAPI only; AR Conversational is a separate application and is not documented here"),
            ("Evidence basis", "ARAPI 1.0.0 source, running desktop build, and the screenshots reproduced in this guide"),
            ("Review date", REVIEW_DATE),
        ],
        [2700, 6660],
        9.6,
    )
    add_callout(
        doc,
        "Client-safe scope",
        "The examples use mock or placeholder values. Never copy production API keys, access tokens, customer data, or confidential response bodies into documentation or support tickets.",
        "caution",
    )
    add_heading(doc, "How to use this guide", 2)
    add_list(doc, [
        "Start with Sections 2-4 before operating a new installation.",
        "Use Section 6 as the visual, screen-by-screen walkthrough; every screenshot is from the running ARAPI desktop app.",
        "Use Section 7 when configuring Designer commands and Section 9 when generating or running Bash/curl files.",
        "Use Appendix A to look up every visible control and Appendix B only for approved local integration/diagnostic work.",
    ])
    page_break(doc)

    toc_heading = add_heading(doc, "Contents", 1, "toc")
    toc_heading.paragraph_format.space_after = Pt(10)
    for anchor, label in TOC:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15 if not label.startswith("Appendix") else 0.35)
        p.paragraph_format.space_after = Pt(4)
        add_internal_link(p, label, anchor, size=10.5, bold=not label.startswith("Appendix"))
    add_callout(doc, "Navigation", "In Word, hold Ctrl while clicking an entry. In most PDF viewers, click normally. Each major section includes a Back to contents link.", "note")
    page_break(doc)


def build_about(doc: Document) -> None:
    add_heading(doc, "1. About ARAPI", 1, "about", True)
    add_para(doc, "ARAPI is a Windows desktop application for catalog-driven API testing. It imports real OpenAPI/Swagger specifications, presents their endpoints in a banking taxonomy, lets users assemble ordered BotJobs, and exports Bash/curl files for command-line execution.")
    add_para(doc, "The user interface is local and communicates with a localhost C# backend. No login is required inside the desktop app; workstation access, file permissions, network controls, and upstream API authentication remain the customer's responsibility.")
    add_heading(doc, "Core concepts", 2)
    for label, text in (
        ("Catalog", "The imported, searchable list of real API methods and paths."),
        ("BotJob", "A saved ordered workflow made of commands and variables."),
        ("Environment", "A named target base URL plus optional default headers."),
        ("Run", "A saved local execution/audit record with step rows."),
        ("Bash export", "A generated .sh file containing enabled API_CALL commands as curl invocations."),
    ):
        add_definition(doc, label, text)
    add_heading(doc, "Recommended operating sequence", 2)
    add_list(doc, [
        "Import specifications and confirm the catalog.",
        "Create an approved environment or use the built-in mock context.",
        "Create/review the BotJob and Save it.",
        "Generate a Bash script into a controlled directory and review it.",
        "Set required variables, run the script, and retain target-side evidence.",
        "Use Execute Tests and Reports for the ARAPI-local audit and exports.",
    ], numbered=True)


def build_release_notes(doc: Document) -> None:
    add_heading(doc, "2. Read this before using the current release", 1, "release_notes", True)
    add_callout(doc, "Version basis", f"This manual describes the observed ARAPI {VERSION} desktop build as reviewed on {REVIEW_DATE}. Screen counts and demo data will differ by workstation.", "note")
    add_heading(doc, "Current-build qualifications", 2)
    add_list(doc, [
        "Test Cases is a reserved workspace: New test case is disabled in version 1.0.0.",
        "IF, ELSE, LOOP, FOR_EACH, READ_CSV, READ_EXCEL, AI_GENERATE_DATA, and CALL_COMPONENT are preview palette items. The Designer states that they are not executed and are recorded as skipped.",
        "BotJob-specific Bash generation includes enabled API_CALL commands only. Assertions, extraction, waits, and control/data commands are not emitted into the .sh file.",
        "Bot Builder may create a synthetic, mock-only endpoint/data record when no imported endpoint matches strongly enough. Review its evidence and never treat a synthetic record as an imported production contract.",
        "The observed desktop C# build uses the Execute Tests page as a local run/audit view. A green row is not independent proof that a real downstream banking system changed state; use the Bash execution path plus target-side logs/data for acceptance.",
        "Mock Start/Stop controls the in-app state/log in the observed build. Do not assume port 8855 is an externally callable HTTP listener unless the deployment administrator has enabled and verified one.",
        "The AI provider Test control is a local configuration indicator in the observed build. Confirm provider connectivity with a real Bot Builder request.",
    ])
    add_callout(doc, "Production safeguard", "Before pointing any environment at a non-mock system, obtain authorization, review the generated curl file, restrict credentials to least privilege, and agree on expected side effects and rollback.", "risk")


def build_install(doc: Document) -> None:
    add_heading(doc, "3. Installation and first launch", 1, "install", True)
    add_heading(doc, "Install or unpack", 2)
    add_list(doc, [
        "Use the ARAPI installer supplied by your organization, or unpack the approved portable ARAPI bundle into a user-writable folder.",
        "Keep the executable, packaged backend, and any supplied data directory together. Do not move individual sidecar files out of the distribution.",
        "Launch ARAPI. The app opens a single resizable desktop window; no web login is required.",
    ], numbered=True)
    add_heading(doc, "Verify first launch", 2)
    add_list(doc, [
        "Confirm the top-right indicator becomes sidecar online.",
        "Open Home and confirm Engine status is online.",
        "Open Settings and confirm Sidecar port 8787 is displayed.",
        "If a portable distribution includes a data folder next to ARAPI.exe, preserve and back up that folder before upgrades. Other distributions use an application-data location selected by the desktop shell.",
    ])
    add_callout(doc, "Local-only service", "The backend is intended for localhost use. Do not publish port 8787 or 8855 on a public network interface.", "caution")


def build_quick_start(doc: Document) -> None:
    add_heading(doc, "4. End-to-end quick start", 1, "quick_start", True)
    steps = [
        "Confirm sidecar online on Home.",
        "Open Import APIs. Use Upload Files for selected files/folders, or Folder (Desktop) for an existing local specification directory. Import .yaml, .yml, or .json files.",
        "Open API Catalog. Search for a known method/path and confirm the summary and mapping status.",
        "Open Environments. For a real approved target, create a named environment with its base URL and non-secret defaults. Keep credentials in variables where possible.",
        "Open BotJob Designer. Create/select a BotJob, add variables, add and configure commands, order them, and Save.",
        "Open Scripts (or select the BotJob in Execute Tests and choose Create Scripts). Select the environment and Browse to Documents\\ARAPI\\Scripts.",
        "Select Create Bash Script. Review the timestamped .sh file; set every required environment variable; run with Bash and curl.",
        "Use Execute Tests for the ARAPI-local run audit, then Reports for HTML/CSV or whole-catalog Postman/Bash downloads.",
    ]
    add_list(doc, steps, numbered=True)
    add_callout(doc, "Best output directory", "Use a dedicated user-writable directory such as %USERPROFILE%\\Documents\\ARAPI\\Scripts. Do not save generated scripts inside the application installation folder.", "success")


def build_interface(doc: Document) -> None:
    add_heading(doc, "5. Interface fundamentals", 1, "interface", True)
    add_heading(doc, "Global header", 2)
    add_definition(doc, "Page name", "The left side of the top bar repeats the current workspace.")
    add_definition(doc, "Employee / Client", "Changes shared conversation-mode context where a mode-aware experience is used. It does not change catalog contents, environment definitions, or BotJob command order.")
    add_definition(doc, "sidecar status", "Green/online means the React interface can reach the local backend. Resolve offline before importing, saving, or running.")
    add_heading(doc, "Sidebar", 2)
    add_table(doc, ["Destination", "Purpose"], NAV_ITEMS, [2600, 6760], 9.4)
    add_heading(doc, "Common visual states", 2)
    add_list(doc, [
        "Blue primary button: the main action for the page.",
        "Dimmed button or palette item: required selection/data is missing, or the feature is reserved.",
        "Green badge: passed, active default, running, or online depending on context.",
        "Orange/red text: warning, failure, missing mapping, or validation error that requires review.",
        "Chevron: expand/collapse a provider, command, action card, or result row.",
    ])
    add_figure(doc, *SCREENS[0])


def build_screens(doc: Document) -> None:
    add_heading(doc, "6. Screen-by-screen guide", 1, "screens", True)
    add_para(doc, "The following pages show every main ARAPI screen and the important interaction states captured from the running desktop application. Values are demonstration data; workflows and controls are the focus.")
    for index, screen in enumerate(SCREENS):
        if index == 0:
            continue
        page_break(doc)
        filename, title, alt, note = screen
        add_heading(doc, f"6.{index + 1} {title}", 2)
        add_figure(doc, filename, title, alt, note)
        if filename == "02-import-apis.png":
            add_callout(doc, "Importer scope", "The recursive scan ignores node_modules, dist, build, .git, bin, and obj. Only endpoints present in parsed specifications enter the catalog.", "note")
        elif filename == "14-import-upload.png":
            add_callout(doc, "Supported files", ".yaml, .yml, and .json. Review the import summary and each reported failure before continuing.", "note")
        elif filename == "03-api-catalog.png":
            add_callout(doc, "Single source of selection", "Designer API_CALL controls can select only endpoint IDs currently returned by the catalog.", "note")
        elif filename == "05-test-cases.png":
            add_callout(doc, "Release status", "New test case is disabled in 1.0.0. Use BotJob Designer for available workflow authoring.", "caution")
        elif filename == "24-bot-builder-action.png":
            add_callout(doc, "Evidence rule", "An action card is evidence of what ARAPI matched or created. Confirm whether the endpoint is imported or synthetic before relying on it.", "caution")
        elif filename == "26-designer-command-config.png":
            add_code_block(doc, '${token}\n{"Authorization": "Bearer ${token}"}\n{"customerId": "${customerId}"}')
        elif filename == "16-environment-form.png":
            add_code_block(doc, "Authorization: Bearer ${token}\nX-Tenant-ID: acme")
        elif filename == "17-scripts-selected.png":
            add_callout(doc, "Recommended directory", "%USERPROFILE%\\Documents\\ARAPI\\Scripts", "success")
        elif filename == "19-execute-results.png":
            add_callout(doc, "Evidence", "Use target-side logs/data for live-system acceptance; a local ARAPI status alone is insufficient.", "caution")
        elif filename == "21-mock-running-log.png":
            add_callout(doc, "Current build", "The screenshot proves ARAPI state/log changes. It does not by itself prove that an external listener is bound to port 8855.", "caution")
        elif filename == "12-reports.png":
            add_callout(doc, "Export distinction", "Run HTML/CSV uses one selected run. Postman and Bash/curl export the whole catalog and default to a localhost base URL unless an integration caller supplies one.", "note")
        elif filename == "22-settings-provider-form.png":
            add_callout(doc, "Credential handling", "The key field is masked and provider-list responses omit key values. Still protect the local ARAPI data store as sensitive and use non-production credentials where possible.", "risk")


def build_commands(doc: Document) -> None:
    add_heading(doc, "7. BotJob command reference", 1, "commands", True)
    add_para(doc, "Commands execute in canvas order when enabled. Click a palette item to add it, use the grip to reorder, the checkbox to enable/disable, the chevron to edit, and x to delete. Save after every material change.")
    add_callout(doc, "Bash boundary", "Only enabled API_CALL commands are exported to the BotJob-specific .sh file. The other command descriptions below apply to the Designer/runtime model and local run audit.", "caution")
    current_group = None
    group_number = 0
    command_number = 0
    for command in COMMANDS:
        if command["group"] != current_group:
            current_group = command["group"]
            group_number += 1
            command_number = 0
            add_heading(doc, f"7.{group_number} {current_group} commands", 2)
        command_number += 1
        add_heading(doc, f"{command['label']} ({command['name']})", 3)
        status_kind = "caution" if command["status"] in ("Preview", "Verify in deployment") else "success"
        add_callout(doc, "Availability", command["status"], status_kind)
        add_definition(doc, "Purpose", command["purpose"])
        add_definition(doc, "Configuration", command["configuration"])
        add_definition(doc, "Current release", command["current"])
    add_heading(doc, "Command design rules", 2)
    add_list(doc, [
        "Place API_CALL before assertions or extraction that depend on its response.",
        "Use catalog endpoint selection; do not paste an unregistered path into descriptive fields.",
        "Use ${name} consistently and declare variables with meaningful names.",
        "Keep secret initial values empty. Mark the variable secret so Bash requires an environment value.",
        "Disable rather than delete a command when temporarily excluding it for diagnosis.",
        "Do not rely on preview commands in client acceptance criteria.",
    ])


def build_variables_env(doc: Document) -> None:
    add_heading(doc, "8. Variables and environments", 1, "variables_env", True)
    add_heading(doc, "Variables", 2)
    add_para(doc, "A BotJob variable has a name, optional initial value, and secret flag. Reference it in supported command text with ${name}.")
    add_table(
        doc,
        ["Element", "Guidance"],
        [
            ("Name", "Use letters, digits, and underscores; prefer clear business names such as customerId or accessToken."),
            ("Initial value", "Use only non-sensitive defaults. Leave secret values empty."),
            ("secret", "For Bash export, requires the normalized shell variable from the environment."),
            ("Normalization", "Bash names become uppercase; non-alphanumeric characters become underscores; a leading digit receives ARAPI_."),
            ("Unknown token", "The exporter treats the normalized token as required and stops before curl if it is unset."),
        ],
        [2700, 6660],
        9.3,
    )
    add_code_block(doc, "customer-id  -> CUSTOMER_ID\n2fa-token    -> ARAPI_2FA_TOKEN\n${accountId} -> ${ACCOUNTID} (normalized by the generated script)")
    add_heading(doc, "Environments", 2)
    add_para(doc, "An environment combines a base URL, optional description, and default headers. One environment can be marked default; the built-in Mock Server environment always exists and cannot be deleted.")
    add_list(doc, [
        "Name and Base URL are mandatory.",
        "Enter default headers one per line as Key: Value.",
        "Command headers override environment headers with the same name in Bash generation.",
        "Sensitive header names without a ${token} are replaced by required Bash environment variables rather than embedded in the generated file.",
        "Environment cards display saved headers in the current UI. Do not save literal production secrets as environment headers.",
    ])
    add_callout(doc, "Safe pattern", "Store Authorization: Bearer ${token}, declare token as a secret variable with no initial value, and supply TOKEN only to the Bash process or approved CI secret store.", "success")


def build_bash(doc: Document) -> None:
    add_heading(doc, "9. Bash/curl script generation and execution", 1, "bash", True)
    add_heading(doc, "Generate from Scripts", 2)
    add_list(doc, [
        "Select the saved BotJob.",
        "Select the environment and verify the displayed base URL.",
        "Select Browse and choose a dedicated folder such as Documents\\ARAPI\\Scripts.",
        "Select Create Bash Script. ARAPI writes a timestamped file such as check-account-balances-20260826-181500.sh.",
        "Read the complete file before running it, especially for POST, PUT, PATCH, or DELETE calls.",
    ], numbered=True)
    add_heading(doc, "Generate from Execute Tests", 2)
    add_para(doc, "After selecting a BotJob, the conditional Create Scripts button opens Scripts with the current BotJob and environment preselected. You still choose the save directory and select Create Bash Script.")
    add_heading(doc, "What the exporter includes", 2)
    add_list(doc, [
        "Enabled API_CALL commands in their saved order.",
        "The selected environment base URL and merged environment/command headers.",
        "JSON request bodies and ${variable} interpolation.",
        "Path parameters such as {accountId} converted to required shell variables.",
        "Required-variable guards for secret variables, sensitive headers, path parameters, and unknown tokens.",
        "A failure counter: each failed curl is reported, remaining calls continue, and the script exits 1 if any call failed.",
    ])
    add_callout(doc, "Not included", "Assertions, extraction, Set Variable steps, waits, stop markers, preview control flow, data-file commands, and component calls are not emitted into the BotJob-specific Bash file.", "caution")
    add_heading(doc, "Generated Bash statement reference", 2)
    statements = [
        ("#!/usr/bin/env bash", "Runs the file with Bash found in the environment."),
        ("# Generated by ARAPI...", "Documents the source BotJob/environment; comments do not execute."),
        ("set -uo pipefail", "Treats unset variables as errors and propagates pipeline failures. The exporter intentionally handles curl failures itself."),
        ("DEFAULT_BASE_URL='...'", "Stores the selected environment URL as the default."),
        ("BASE_URL=\"${BASE_URL:-$DEFAULT_BASE_URL}\"", "Allows a runtime BASE_URL override without editing the file."),
        ("failures=0", "Initializes the curl failure counter."),
        (": \"${TOKEN:?Set TOKEN...}\"", "Stops immediately when a required value is unset or empty."),
        ("DEFAULT_NAME='value'", "Stores a non-secret BotJob initial value."),
        ("NAME=\"${NAME:-$DEFAULT_NAME}\"", "Allows an environment override or uses the non-secret default."),
        ("echo '==> n/N METHOD /path'", "Prints progress before each call."),
        ("if ! curl \\", "Runs curl and enters the failure branch when curl returns non-zero."),
        ("--globoff", "Prevents curl from treating braces/brackets in URLs as glob patterns."),
        ("--fail-with-body", "Returns failure for HTTP 4xx/5xx while retaining the response body."),
        ("--silent --show-error", "Suppresses progress meter but keeps error messages."),
        ("--request METHOD", "Uses the catalog HTTP method."),
        ("--header 'Key: value'", "Adds merged environment and command headers."),
        ("--data 'JSON'", "Sends the optional body and causes application/json to be added when Content-Type was absent."),
        ("\"${BASE_URL}\"'/path'", "Combines runtime/default base URL with the catalog path."),
        ("failures=$((failures + 1))", "Counts a failed curl while allowing the next API call."),
        ("if (( failures > 0 )); then ... exit 1", "Returns a non-zero process result when any API call failed."),
        ("echo 'All API calls completed successfully.'", "Prints only when the failure counter is zero."),
    ]
    add_table(doc, ["Statement", "Meaning"], statements, [3300, 6060], 8.7)
    add_heading(doc, "Run on Windows with Git Bash", 2)
    add_code_block(doc, "cd /c/Users/YOUR_NAME/Documents/ARAPI/Scripts\nexport TOKEN='temporary-secret-value'\nexport ACCOUNT_ID='12345'\nbash ./script-name.sh\necho $?")
    add_heading(doc, "Run from Windows PowerShell when bash is on PATH", 2)
    add_code_block(doc, "$env:TOKEN = 'temporary-secret-value'\n$env:ACCOUNT_ID = '12345'\n$env:BASE_URL = 'https://approved-test-api.example'\nbash ./script-name.sh\n$LASTEXITCODE\nRemove-Item Env:TOKEN, Env:ACCOUNT_ID")
    add_heading(doc, "Run from Windows Command Prompt", 2)
    add_code_block(doc, "cd /d \"%USERPROFILE%\\Documents\\ARAPI\\Scripts\"\nset \"TOKEN=temporary-secret-value\"\nset \"ACCOUNT_ID=12345\"\nbash script-name.sh\necho %ERRORLEVEL%\nset \"TOKEN=\"\nset \"ACCOUNT_ID=\"")
    add_heading(doc, "Run with Windows Subsystem for Linux", 2)
    add_code_block(doc, "cd /mnt/c/Users/YOUR_NAME/Documents/ARAPI/Scripts\nexport TOKEN='temporary-secret-value'\nbash ./script-name.sh\necho $?")
    add_heading(doc, "Pre-run review checklist", 2)
    add_list(doc, [
        "Correct BotJob name and timestamped file.",
        "Correct BASE_URL and approved environment.",
        "No unintended destructive methods or paths.",
        "Headers/bodies contain no embedded secret or real customer data.",
        "All required shell variables are set from an approved source.",
        "curl is available, TLS trust is configured, and network routing is approved.",
        "Expected side effects, monitoring, and rollback are agreed before execution.",
    ])


def build_results(doc: Document) -> None:
    add_heading(doc, "10. Execution results, history, and reports", 1, "results", True)
    add_heading(doc, "Run audit", 2)
    add_para(doc, "Execute Tests stores a run summary and step rows. Summary fields are Status, Target, Steps, Duration, and Run ID. Step rows show command type, status, duration, and expandable details when recorded.")
    add_table(
        doc,
        ["Status", "Interpretation"],
        [
            ("passed", "The local run record marks the step/run successful."),
            ("failed", "An assertion or step reported a failure."),
            ("error", "The step could not complete normally, for example a missing endpoint reference."),
            ("skipped", "The command was disabled, stopped, preview-only, or otherwise not executed."),
            ("running", "The run has not yet reached its final state."),
        ],
        [1900, 7460],
        9.4,
    )
    add_callout(doc, "Acceptance rule", "For a real environment, require target-side logs, database/state evidence, or an approved response capture in addition to the ARAPI run status.", "risk")
    add_heading(doc, "History", 2)
    add_list(doc, [
        "Select History to show saved runs.",
        "Use All BotJobs to filter by workflow.",
        "Select a row to reload the summary and steps; the panel closes automatically.",
    ], numbered=True)
    add_heading(doc, "Reports", 2)
    add_table(
        doc,
        ["Export", "Scope", "Use"],
        [
            ("HTML Report", "Selected run", "Readable summary; can be printed to PDF by the user's browser/viewer."),
            ("CSV Export", "Selected run", "Tabular step results for spreadsheet or BI review."),
            ("Postman Collection", "Whole catalog", "Postman Collection v2.1 JSON; review the default localhost base URL."),
            ("Bash / curl Script", "Whole catalog", "One simple curl per endpoint; unlike Scripts, it is not BotJob-specific."),
        ],
        [2100, 1800, 5460],
        9.1,
    )
    add_callout(doc, "Prefer Scripts for controlled execution", "The Scripts page preserves a selected BotJob's enabled API_CALL order, environment headers, body data, variable guards, and final failure status. Reports Bash/curl is a broad catalog export.", "note")


def build_ai(doc: Document) -> None:
    add_heading(doc, "11. Bot Builder and AI providers", 1, "ai", True)
    add_heading(doc, "Bot Builder", 2)
    add_para(doc, "Bot Builder sends the visible conversation history to the local backend, which can search the catalog, list/create BotJobs, and request mock runs. Action cards expose catalog matches, created jobs, and run summaries.")
    add_list(doc, [
        "Use a suggestion or enter a precise business flow.",
        "Review catalog-search evidence before accepting a proposed workflow.",
        "If the result says a synthetic mock endpoint/data record was created, keep it out of production acceptance and exports intended to represent real contracts.",
        "Open Designer and review the created job, endpoint selections, body, headers, variables, and command order before Save or execution.",
        "Do not paste credentials, customer identifiers, account data, or confidential payloads into the conversation.",
    ])
    add_heading(doc, "Supported provider records", 2)
    add_table(
        doc,
        ["Provider", "Default model shown by ARAPI", "Base URL field"],
        [
            ("OpenAI", "gpt-4o-mini", "Not shown"),
            ("Anthropic Claude", "claude-3-5-haiku-20241022", "Not shown"),
            ("Google Gemini", "gemini-1.5-flash", "Not shown"),
            ("Azure OpenAI", "gpt-4o-mini", "Required by deployment"),
            ("Ollama (local)", "llama3.2", "Shown; key not required for local Ollama"),
            ("Together.ai", "meta-llama/Llama-3-8b-chat-hf", "Not shown"),
            ("Custom OpenAI-compatible", "gpt-4o-mini", "Required by deployment"),
        ],
        [2500, 4100, 2760],
        8.9,
    )
    add_heading(doc, "Configure a provider", 2)
    add_list(doc, [
        "Expand the provider card.",
        "Enter or confirm the model and base URL where shown.",
        "Enter the API key and select Save. The field is masked; a blank key on later saves preserves the existing key.",
        "Select SET DEFAULT on a configured provider. If exactly one key exists, ARAPI may promote it automatically.",
        "Select Test, then make an actual Bot Builder request to verify end-to-end connectivity.",
    ], numbered=True)
    add_callout(doc, "Local data sensitivity", "Provider-list responses omit key values, but the workstation and ARAPI data directory must still be protected as credential-bearing assets. Use scoped, revocable, non-production keys for testing.", "risk")


def build_security(doc: Document) -> None:
    add_heading(doc, "12. Security and operational practices", 1, "security", True)
    add_heading(doc, "Before use", 2)
    add_list(doc, [
        "Use an organization-managed Windows account and approved ARAPI distribution.",
        "Keep the local backend bound to localhost; do not expose service ports publicly.",
        "Back up the correct ARAPI data directory before upgrades or data-cleanup work.",
        "Import specifications only from trusted sources and review import failures.",
        "Create separate mock/dev/UAT/production environments and label them clearly.",
    ])
    add_heading(doc, "Credentials and data", 2)
    add_list(doc, [
        "Never store literal secrets in BotJob initial values, non-secret headers, descriptions, screenshots, or client documents.",
        "Use secret BotJob variables and short-lived environment variables or an approved CI secret store.",
        "Generated scripts and exported reports may contain paths, bodies, headers, response details, and business metadata; classify and retain them accordingly.",
        "Review scripts after generation and again before every non-mock run.",
        "Clear shell environment variables and delete temporary artifacts according to policy after use.",
    ])
    add_heading(doc, "Change control", 2)
    add_list(doc, [
        "Record ARAPI version, BotJob name, environment, script filename/hash, operator, start/end time, and target approval.",
        "Use a least-privilege API identity with a narrow expiration window.",
        "For write operations, document expected records and rollback/compensation before execution.",
        "Preserve HTML/CSV and target-side evidence in the approved audit location; do not rely on screenshots alone.",
    ])
    add_callout(doc, "No-login does not mean no security", "ARAPI's desktop UI has no internal login. Operating-system access, file ACLs, secret management, API authorization, and audit retention provide the security boundary.", "caution")


def build_troubleshooting(doc: Document) -> None:
    add_heading(doc, "13. Troubleshooting", 1, "troubleshooting", True)
    add_para(doc, "Use the exact displayed error, affected BotJob/environment, ARAPI version, and reproduction steps when escalating. Remove secrets and customer data from attachments.")
    for issue, resolution in TROUBLESHOOTING:
        add_heading(doc, issue, 2)
        add_para(doc, resolution)
    add_heading(doc, "Diagnostic information safe to collect", 2)
    add_list(doc, [
        "ARAPI version and whether installer or portable distribution is used.",
        "Home engine status and Settings local-service port values.",
        "Specification filename/type and import counts (not confidential contents unless approved).",
        "BotJob name, command types/order, and whether the failing command is enabled.",
        "Environment name and sanitized base URL; never send tokens or full Authorization headers.",
        "Script filename, exit code, and sanitized curl error text.",
        "Run ID and report with sensitive request/response content removed.",
    ])


def build_status(doc: Document) -> None:
    add_heading(doc, "14. Feature status matrix", 1, "status", True)
    rows = [
        ("Import YAML/JSON files/folders", "Available", "Recursive upload/desktop folder workflows."),
        ("Catalog search and mapping display", "Available", "Mapping is shown; no mapping editor button on the Catalog page."),
        ("25-domain banking taxonomy", "Available", "Read-only cards in current UI."),
        ("Test Case creation", "Reserved", "New test case is disabled."),
        ("Bot Builder catalog/search/create/run actions", "Available with qualifications", "May use synthetic mock endpoint/data when no strong catalog match."),
        ("Designer core commands", "Configurable", "API, variables, assertions, extract, wait; verify runtime acceptance per deployment."),
        ("Designer control/data preview commands", "Preview", "IF/ELSE/loops/read files/AI data/component are recorded skipped."),
        ("Environment CRUD", "Available", "Built-in Mock cannot be deleted."),
        ("BotJob-specific Bash/curl", "Available", "Enabled API_CALL commands only; desktop exact-directory picker."),
        ("Execute Tests audit/history", "Available with qualification", "Local run evidence; correlate real-system effects externally."),
        ("Mock dashboard", "State/log available", "External port listener must be deployment-verified."),
        ("HTML/CSV run reports", "Available", "Requires a saved run."),
        ("Postman/whole-catalog Bash", "Available", "Defaults to localhost base URL via UI download."),
        ("AI provider records/default selection", "Available", "Protect local data store as sensitive."),
        ("AI provider Test", "Configuration indicator", "Use a real Bot Builder request for end-to-end proof."),
    ]
    add_table(doc, ["Capability", "Status", "Qualification"], rows, [3400, 2100, 3860], 8.7)


def build_appendices(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "Appendix A. Complete control index", 1, "controls", True)
    add_para(doc, "This index covers every visible button, selector, disclosure control, status control, and important conditional action in the ARAPI 1.0.0 UI.")
    add_table(doc, ["Page", "Control", "Available when", "Result"], CONTROL_INDEX, [1512, 2304, 1800, 3744], 8.1)

    page_break(doc)
    add_heading(doc, "Appendix B. Local API reference", 1, "api", True)
    add_callout(doc, "Administrator/developer reference", "These JSON routes are localhost-only implementation interfaces. Normal client operation should use the ARAPI UI and generated Bash files. Never publish this API publicly.", "caution")
    add_table(doc, ["Method", "Route", "Purpose"], API_ROUTES, [900, 4050, 4410], 8.2)

    page_break(doc)
    add_heading(doc, "Appendix C. Glossary", 1, "glossary", True)
    add_table(doc, ["Term", "Meaning"], GLOSSARY, [2500, 6860], 9.4)

    page_break(doc)
    add_heading(doc, "Appendix D. Client acceptance checklist", 1, "handoff", True)
    add_heading(doc, "Application", 2)
    add_list(doc, [
        "ARAPI version and distribution type are recorded.",
        "sidecar online appears after a clean launch.",
        "Approved backup/restore location is documented and tested.",
        "Local ports are not exposed publicly.",
    ])
    add_heading(doc, "Catalog and workflow", 2)
    add_list(doc, [
        "A known-good specification imports with expected endpoint count.",
        "A sample catalog endpoint is searchable and correctly classified or marked unmapped.",
        "A BotJob can be created, configured, reordered, and saved.",
        "Preview commands are excluded from acceptance criteria.",
    ])
    add_heading(doc, "Bash execution", 2)
    add_list(doc, [
        "The Scripts page saves to Documents\\ARAPI\\Scripts or an approved equivalent.",
        "The generated file contains only intended enabled API_CALL commands.",
        "Secrets appear as required environment variables, not literal values.",
        "A mock/dev script returns the expected process exit code and target-side evidence.",
    ])
    add_heading(doc, "Audit and AI", 2)
    add_list(doc, [
        "Execution History reloads a saved run and Reports downloads HTML/CSV.",
        "Whole-catalog Postman/Bash files are reviewed for localhost base URL before use.",
        "AI provider configuration is verified with a real Bot Builder request.",
        "Operators understand synthetic endpoint labels and do not treat them as imported production APIs.",
    ])
    add_callout(doc, "Sign-off", "Client acceptance should record approver, date, environment, BotJob/script identifier, evidence location, exceptions, and remediation owner.", "success")
    bottom = doc.add_paragraph()
    bookmark_paragraph(bottom, "doc_bottom")
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_internal_link(bottom, "Return to contents", "toc", size=9.5)


def build() -> Path:
    doc = Document()
    configure_document(doc)
    install_numbering(doc)
    build_cover(doc)
    build_front_matter(doc)
    build_about(doc)
    build_release_notes(doc)
    build_install(doc)
    build_quick_start(doc)
    build_interface(doc)
    build_screens(doc)
    build_commands(doc)
    build_variables_env(doc)
    build_bash(doc)
    build_results(doc)
    build_ai(doc)
    build_security(doc)
    build_troubleshooting(doc)
    build_status(doc)
    build_appendices(doc)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
